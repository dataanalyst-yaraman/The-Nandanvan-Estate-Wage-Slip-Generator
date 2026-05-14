"""
Weekly Wage Slip Generator
Streamlit application that fetches attendance data from Google Sheets
and generates DOCX wage slips for each employee for a selected week
(Thursday–Wednesday).
"""

import os
import io
import re
from datetime import date, timedelta, datetime

import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
load_dotenv()

HEADER_IMAGE_PATH = os.path.join(os.path.dirname(__file__), "wage_slip_header_image.png")
WAGES_SHEET_NAME = "Wages"

# ---------------------------------------------------------------------------
# Google Sheets connection (cached)
# ---------------------------------------------------------------------------

def _get_secret(key: str, default: str = "") -> str:
    """Read from st.secrets (Streamlit Cloud) first, then fall back to .env."""
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError):
        return os.getenv(key, default)


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    """Authenticate and return a gspread client using secrets / .env credentials."""
    private_key = _get_secret("GCP_PRIVATE_KEY")
    # .env files may double-escape newlines
    if "\\n" in private_key and "\n" not in private_key.replace("\\n", ""):
        private_key = private_key.replace("\\n", "\n")

    client_email = _get_secret("GCP_CLIENT_EMAIL")
    creds_dict = {
        "type": _get_secret("GCP_TYPE"),
        "project_id": _get_secret("GCP_PROJECT_ID"),
        "private_key_id": _get_secret("GCP_PRIVATE_KEY_ID"),
        "private_key": private_key,
        "client_email": client_email,
        "client_id": _get_secret("GCP_CLIENT_ID"),
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
        "client_x509_cert_url": (
            f"https://www.googleapis.com/robot/v1/metadata/x509/{client_email}"
        ),
    }
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets.readonly",
        "https://www.googleapis.com/auth/drive.readonly",
    ]
    credentials = Credentials.from_service_account_info(creds_dict, scopes=scopes)
    return gspread.authorize(credentials)


def fetch_wages_data():
    """Fetch all rows from the 'Wages' worksheet."""
    client = get_gspread_client()
    sheet_url = _get_secret("GCP_SHEET_URL")
    spreadsheet = client.open_by_url(sheet_url)
    ws = spreadsheet.worksheet(WAGES_SHEET_NAME)
    return ws.get_all_values()


# ---------------------------------------------------------------------------
# Date helpers
# ---------------------------------------------------------------------------

def parse_date(date_str: str) -> date | None:
    """Parse 'M/D/YYYY' or 'MM/DD/YYYY' date strings from the sheet."""
    if not date_str:
        return None
    try:
        return datetime.strptime(date_str.strip(), "%m/%d/%Y").date()
    except ValueError:
        try:
            return datetime.strptime(date_str.strip(), "%d/%m/%Y").date()
        except ValueError:
            return None


# Day-of-week constants for the selectbox
DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]


def get_week_bounds(any_date: date, week_start_weekday: int = 3) -> tuple[date, date]:
    """
    Return (week_start, week_end) for the week containing *any_date*.
    week_start_weekday: 0=Mon … 6=Sun (default 3=Thursday).
    The week is always 7 days: start_day → start_day + 6.
    """
    day_of_week = any_date.weekday()
    days_since_start = (day_of_week - week_start_weekday) % 7
    week_start = any_date - timedelta(days=days_since_start)
    week_end = week_start + timedelta(days=6)
    return week_start, week_end


def get_payment_date(week_end: date) -> date:
    """Payment date = day after the week end date."""
    return week_end + timedelta(days=1)


# ---------------------------------------------------------------------------
# Currency helpers
# ---------------------------------------------------------------------------

def parse_currency(value: str) -> float:
    """Strip ₹ and commas, return float.  Returns 0.0 for blanks / errors."""
    if not value:
        return 0.0
    cleaned = re.sub(r"[₹,\s]", "", value)
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def fmt_inr(amount: float) -> str:
    """Format a number as ₹ X,XXX.XX"""
    return f"₹{amount:,.2f}"


# ---------------------------------------------------------------------------
# Aggregate data per employee for a given week
# ---------------------------------------------------------------------------

def aggregate_week(raw_rows: list[list[str]], thu: date, wed: date) -> list[dict]:
    """
    Given all raw rows (including header row) and a week range,
    aggregate per employee and return a list of dicts ready for the slip.
    """
    # Header: Date | Day | Employee Code | Name | Department | Designation
    #         | Attendance | Wages | Overtime Hours | Overtime Wages | Total Wages
    header = raw_rows[0]
    data_rows = raw_rows[1:]

    # Group by Employee Code
    employees: dict[str, dict] = {}
    for row in data_rows:
        row_date = parse_date(row[0])
        if row_date is None:
            continue
        if not (thu <= row_date <= wed):
            continue

        emp_code = row[2].strip()
        attendance = row[6].strip()

        if emp_code not in employees:
            employees[emp_code] = {
                "code": emp_code,
                "name": row[3].strip(),
                "department": row[4].strip(),
                "designation": row[5].strip(),
                "working_days": 0,
                "overtime_hours": 0.0,
                "total_wages": 0.0,
                "total_overtime_wages": 0.0,
                "deductions": 0.0,
            }

        emp = employees[emp_code]
        if attendance.lower() == "present":
            emp["working_days"] += 1

        emp["total_wages"] += parse_currency(row[7])       # Daily wages
        ot_hours = row[8].strip()
        emp["overtime_hours"] += float(ot_hours) if ot_hours else 0.0
        emp["total_overtime_wages"] += parse_currency(row[9])

    # Sort by employee code
    return sorted(employees.values(), key=lambda e: e["code"])


# ---------------------------------------------------------------------------
# DOCX generation  (matches the template layout)
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    """
    Set cell border.
    Usage: _set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs["val"]}" '
            f'w:sz="{attrs["sz"]}" w:space="0" w:color="{attrs["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_row_height(row, height_emu: int):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{height_emu}" w:hRule="exact"/>'
    )
    trPr.append(trHeight)


def _merge_cells(table, row_idx, col_start, col_end):
    """Merge cells horizontally in a row."""
    cell_start = table.cell(row_idx, col_start)
    cell_end = table.cell(row_idx, col_end)
    cell_start.merge(cell_end)


def _write_cell(cell, text, bold=False, size=10, alignment=None, font_name="Calibri",
                color=None, vertical_alignment=None):
    """Write styled text to a table cell (clears existing content first)."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()
    paragraph = cell.paragraphs[0]
    if alignment:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = color

    # Vertical alignment
    if vertical_alignment:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = parse_xml(
            f'<w:vAlign {nsdecls("w")} w:val="{vertical_alignment}"/>'
        )
        tcPr.append(vAlign)


def _apply_table_borders(table):
    """Apply uniform borders to all cells in a table."""
    border_attrs = {"sz": 4, "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top=border_attrs,
                bottom=border_attrs,
                start=border_attrs,
                end=border_attrs,
            )


def get_all_weeks_in_range(
    start_dt: date, end_dt: date, week_start_weekday: int = 3
) -> list[tuple[date, date, date]]:
    """
    Return a list of (week_start, week_end, payment_date) for every
    wage week that overlaps with [start_dt, end_dt].
    """
    first_start, _ = get_week_bounds(start_dt, week_start_weekday)
    weeks = []
    ws = first_start
    while ws <= end_dt:
        we = ws + timedelta(days=6)
        pay = get_payment_date(we)
        weeks.append((ws, we, pay))
        ws += timedelta(days=7)
    return weeks


def _add_wage_slip_page(doc, emp: dict, week_start: date, week_end: date,
                        payment_date: date):
    """Add a single wage-slip page for one employee / one week."""
    wage_period_str = (
        f"{week_start.strftime('%d %b %Y')} – {week_end.strftime('%d %b %Y')}"
    )
    payment_date_str = payment_date.strftime("%d %b %Y")

    # ── Header image ──
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(4)
    if os.path.exists(HEADER_IMAGE_PATH):
        run_img = p_img.add_run()
        run_img.add_picture(HEADER_IMAGE_PATH, width=Inches(4.8))

    # ── Title ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run("WEEKLY WAGE SLIP")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.name = "Calibri"
    run_title.font.color.rgb = RGBColor(0x2E, 0x4A, 0x1F)

    # ── Main table (13 rows × 3 columns) ──
    # Rows 0-10: col 0 = label, cols 1+2 merged = value
    # Row 11: 3 separate signature labels
    # Row 12: 3 separate blank signature spaces
    table = doc.add_table(rows=13, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    col_w0 = Cm(6.0)
    col_w1 = Cm(6.0)
    col_w2 = Cm(6.0)
    for row in table.rows:
        row.cells[0].width = col_w0
        row.cells[1].width = col_w1
        row.cells[2].width = col_w2

    # Calculate weekly amount paid
    weekly_base = emp["total_wages"]
    weekly_ot = emp["total_overtime_wages"]
    deductions = emp["deductions"]
    weekly_amount = weekly_base + weekly_ot - deductions

    # Overtime display text
    ot_hours = emp["overtime_hours"]
    overtime_text = f"{ot_hours:.0f} hrs" if ot_hours > 0 else "Nil"
    if weekly_ot > 0:
        overtime_text += f"  ({fmt_inr(weekly_ot)})"

    # Row data: (label, value)
    rows_data = [
        ("Company Name", "The Nandanvan Estate"),
        ("Department", emp["department"]),
        ("Designation", emp["designation"] or "—"),
        ("Employee Name", emp["name"]),
        ("No. of Working Days", str(emp["working_days"])),
        ("Overtime", overtime_text),
        ("Deductions", fmt_inr(deductions) if deductions > 0 else "Nil"),
        ("Wage Period", wage_period_str),
        ("Date of Payment", payment_date_str),
        ("Weekly Amount Paid (INR)", fmt_inr(weekly_amount)),
        ("Payment Mode", "Bank Transfer"),
    ]

    # Fill rows 0–10: merge cols 1+2 for the value column
    for r_idx, (label, value) in enumerate(rows_data):
        # Merge value cells (cols 1 and 2)
        table.cell(r_idx, 1).merge(table.cell(r_idx, 2))

        label_cell = table.cell(r_idx, 0)
        value_cell = table.cell(r_idx, 1)

        _write_cell(label_cell, label, bold=True, size=10, font_name="Calibri",
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, vertical_alignment="center")
        _write_cell(value_cell, value, bold=False, size=10, font_name="Calibri",
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, vertical_alignment="center")

        # Style the header row (Company Name)
        if r_idx == 0:
            _set_cell_shading(label_cell, "2E4A1F")
            _write_cell(label_cell, label, bold=True, size=10, font_name="Calibri",
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        color=RGBColor(0xFF, 0xFF, 0xFF),
                        vertical_alignment="center")
            _set_cell_shading(value_cell, "F5F0E8")

        # Style the amount row
        if r_idx == 9:
            _set_cell_shading(label_cell, "2E4A1F")
            _write_cell(label_cell, label, bold=True, size=10, font_name="Calibri",
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        color=RGBColor(0xFF, 0xFF, 0xFF),
                        vertical_alignment="center")
            _set_cell_shading(value_cell, "F5F0E8")
            _write_cell(value_cell, value, bold=True, size=11, font_name="Calibri",
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        vertical_alignment="center")

        # Set consistent row height
        _set_row_height(table.rows[r_idx], 340)

    # Row 11: Three separate signature labels
    sig_row = table.rows[11]
    _set_row_height(sig_row, 240)
    _write_cell(sig_row.cells[0], "Prepared By", bold=True, size=9,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_alignment="center")
    _write_cell(sig_row.cells[1], "Approved By", bold=True, size=9,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_alignment="center")
    _write_cell(sig_row.cells[2], "Employee Signature", bold=True, size=9,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_alignment="center")

    # Row 12: Three separate blank signature spaces
    blank_row = table.rows[12]
    _set_row_height(blank_row, 600)
    _write_cell(blank_row.cells[0], "", size=9)
    _write_cell(blank_row.cells[1], "", size=9)
    _write_cell(blank_row.cells[2], "", size=9)

    # Apply borders
    _apply_table_borders(table)


def create_wage_slip_doc(
    weeks_data: list[tuple[date, date, date, list[dict]]],
) -> io.BytesIO:
    """
    Create a Word document with one wage slip per employee per week.
    weeks_data: list of (week_start, week_end, payment_date, employees)
    Returns the document as a BytesIO buffer.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Emu(7560310)   # match template
    section.page_height = Emu(10692130)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    first_page = True
    for week_start, week_end, payment_date, employees in weeks_data:
        for emp in employees:
            if not first_page:
                doc.add_page_break()
            _add_wage_slip_page(doc, emp, week_start, week_end, payment_date)
            first_page = False

    # Write to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------

def main():
    st.set_page_config(
        page_title="Weekly Wage Slip Generator",
        page_icon="📄",
        layout="centered",
    )

    # ── Custom CSS ──
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global */
    .stApp {
        font-family: 'Inter', sans-serif;
    }
    
    /* Header */
    .main-header {
        background: linear-gradient(135deg, #2E4A1F 0%, #4A7C32 50%, #6B9B4E 100%);
        padding: 2rem 2rem 1.5rem 2rem;
        border-radius: 16px;
        margin-bottom: 1.5rem;
        box-shadow: 0 8px 32px rgba(46, 74, 31, 0.25);
        text-align: center;
    }
    .main-header h1 {
        color: #FFFFFF;
        font-size: 1.8rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: 0.5px;
    }
    .main-header p {
        color: #D4E8C7;
        font-size: 0.95rem;
        margin: 0.5rem 0 0 0;
        font-weight: 300;
    }
    
    /* Cards */
    .info-card {
        background: linear-gradient(135deg, #f8faf6 0%, #eef4e9 100%);
        border: 1px solid #c5d9b8;
        border-radius: 12px;
        padding: 1.2rem 1.5rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    .info-card h4 {
        color: #2E4A1F;
        margin: 0 0 0.3rem 0;
        font-size: 0.85rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }
    .info-card .value {
        color: #1a2e12;
        font-size: 1.25rem;
        font-weight: 700;
    }
    
    /* Employee table */
    .emp-card {
        background: #ffffff;
        border: 1px solid #e2e8d9;
        border-radius: 10px;
        padding: 0.8rem 1rem;
        margin-bottom: 0.5rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
        transition: all 0.2s ease;
    }
    .emp-card:hover {
        border-color: #4A7C32;
        box-shadow: 0 2px 12px rgba(74, 124, 50, 0.15);
        transform: translateY(-1px);
    }
    .emp-name {
        font-weight: 600;
        color: #2E4A1F;
        font-size: 0.95rem;
    }
    .emp-detail {
        color: #6b7c60;
        font-size: 0.8rem;
    }
    .emp-amount {
        font-weight: 700;
        color: #2E4A1F;
        font-size: 1.05rem;
    }
    
    /* Status badges */
    .badge-success {
        display: inline-block;
        background: linear-gradient(135deg, #4A7C32, #6B9B4E);
        color: #fff;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 600;
        letter-spacing: 0.5px;
    }
    .badge-warning {
        display: inline-block;
        background: linear-gradient(135deg, #c79100, #e8a800);
        color: #fff;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 600;
        letter-spacing: 0.5px;
    }
    
    /* Divider */
    .section-divider {
        height: 2px;
        background: linear-gradient(to right, transparent, #c5d9b8, transparent);
        margin: 1rem 0;
        border: none;
    }
    
    /* Footer */
    .footer-text {
        text-align: center;
        color: #8a9a7e;
        font-size: 0.75rem;
        padding: 1rem 0;
        border-top: 1px solid #e2e8d9;
        margin-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Logo ──
    if os.path.exists(HEADER_IMAGE_PATH):
        col_l, col_logo, col_r = st.columns([1, 2, 1])
        with col_logo:
            st.image(HEADER_IMAGE_PATH, use_container_width=True)

    # ── Header ──
    st.markdown("""
    <div class="main-header">
        <h1>📄 Weekly Wage Slip Generator</h1>
        <p>The Nandanvan Estate — Organic Certified Estate</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Date selection ──
    st.markdown("### 📅 Select Date Range")
    st.caption(
        "Pick a start and end date — the app will automatically find all "
        "wage weeks within that range."
    )

    col_start, col_end = st.columns(2)
    with col_start:
        start_date = st.date_input(
            "Start Date",
            value=date.today() - timedelta(days=6),
            format="DD/MM/YYYY",
        )
    with col_end:
        end_date = st.date_input(
            "End Date",
            value=date.today(),
            format="DD/MM/YYYY",
        )

    # ── Week day configuration ──
    st.markdown("#### ⚙️ Week Day Configuration")
    col_ws, col_we = st.columns(2)
    with col_ws:
        week_start_day_name = st.selectbox(
            "Week Starts On",
            options=DAY_NAMES,
            index=3,  # default: Thursday
        )
    week_start_weekday = DAY_NAMES.index(week_start_day_name)  # 0=Mon..6=Sun
    week_end_weekday = (week_start_weekday + 6) % 7
    week_end_day_name = DAY_NAMES[week_end_weekday]

    with col_we:
        st.text_input(
            "Week Ends On",
            value=week_end_day_name,
            disabled=True,
            help="Auto-calculated: 6 days after the start day.",
        )

    # Payment day info
    payment_day_weekday = (week_end_weekday + 1) % 7
    payment_day_name = DAY_NAMES[payment_day_weekday]
    st.markdown(f"""
    <div class="info-card" style="padding:0.6rem 1.2rem;">
        <span style="font-weight:600;color:#2E4A1F;">Payment Day:</span>
        <span style="color:#1a2e12; font-weight:700;"> {payment_day_name}</span>
        <span style="color:#6b7c60; font-size:0.85rem;">
            &nbsp;(day after {week_end_day_name})
        </span>
    </div>
    """, unsafe_allow_html=True)

    if start_date > end_date:
        st.error("⚠️ Start date must be on or before end date.")
        return

    weeks = get_all_weeks_in_range(start_date, end_date, week_start_weekday)

    # Display computed weeks
    st.markdown(f"""
    <div class="info-card">
        <h4>Wage Weeks Found</h4>
        <div class="value">{len(weeks)} week(s)</div>
    </div>
    """, unsafe_allow_html=True)

    for thu, wed, pay in weeks:
        st.markdown(f"""
        <div class="info-card" style="padding:0.6rem 1.2rem; margin-bottom:0.4rem;">
            <span style="font-weight:600;color:#2E4A1F;">
                {thu.strftime('%a, %d %b %Y')} → {wed.strftime('%a, %d %b %Y')}
            </span>
            <span style="color:#6b7c60; font-size:0.85rem;">
                 &nbsp;|&nbsp; Payment: {pay.strftime('%a, %d %b %Y')}
            </span>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ── Fetch data ──
    st.markdown("### 👥 Employee Wage Summary")

    if st.button("🔄 Fetch Data from Google Sheets", type="primary", use_container_width=True):
        st.session_state["raw_data"] = None  # Force refresh

    with st.spinner("Fetching attendance data from Google Sheets..."):
        try:
            if "raw_data" not in st.session_state or st.session_state["raw_data"] is None:
                st.session_state["raw_data"] = fetch_wages_data()
            raw_data = st.session_state["raw_data"]
        except Exception as e:
            st.error(f"❌ Failed to connect to Google Sheets: {e}")
            st.info("Check that the sheet is shared with the service account email.")
            return

    # Aggregate per week
    weeks_data: list[tuple[date, date, date, list[dict]]] = []
    grand_total = 0.0
    total_slips = 0

    for thu, wed, pay in weeks:
        emps = aggregate_week(raw_data, thu, wed)
        if emps:
            weeks_data.append((thu, wed, pay, emps))

    if not weeks_data:
        st.warning(
            f"⚠️ No attendance data found for the selected range "
            f"**{start_date.strftime('%d %b %Y')} – {end_date.strftime('%d %b %Y')}**.\n\n"
            f"Make sure the 'Wages' sheet contains entries within this date range."
        )
        return

    # ── Display per-week summaries ──
    for thu, wed, pay, emps in weeks_data:
        week_label = f"{thu.strftime('%d %b')} – {wed.strftime('%d %b %Y')}"
        st.markdown(f"#### 📆 Week: {week_label}")

        week_total = 0.0
        for emp in emps:
            amt = emp["total_wages"] + emp["total_overtime_wages"] - emp["deductions"]
            week_total += amt
            total_slips += 1

            ot_info = ""
            if emp["overtime_hours"] > 0:
                ot_info = f" • OT: {emp['overtime_hours']:.0f}h ({fmt_inr(emp['total_overtime_wages'])})"

            st.markdown(f"""
            <div class="emp-card">
                <div>
                    <div class="emp-name">{emp['name']}</div>
                    <div class="emp-detail">{emp['department']} • {emp['designation'] or '—'} • {emp['working_days']} day(s){ot_info}</div>
                </div>
                <div class="emp-amount">{fmt_inr(amt)}</div>
            </div>
            """, unsafe_allow_html=True)

        grand_total += week_total
        st.markdown(f"""
        <div class="info-card" style="margin-top:0.3rem;">
            <h4>Week Total</h4>
            <div class="value">{fmt_inr(week_total)}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ── Summary stats ──
    col_s1, col_s2, col_s3 = st.columns(3)
    with col_s1:
        st.markdown(f"""
        <div class="info-card">
            <h4>Total Weeks</h4>
            <div class="value">{len(weeks_data)}</div>
        </div>
        """, unsafe_allow_html=True)
    with col_s2:
        st.markdown(f"""
        <div class="info-card">
            <h4>Total Wage Slips</h4>
            <div class="value">{total_slips}</div>
        </div>
        """, unsafe_allow_html=True)
    with col_s3:
        st.markdown(f"""
        <div class="info-card">
            <h4>Grand Total Payout</h4>
            <div class="value">{fmt_inr(grand_total)}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ── Generate & download ──
    st.markdown("### 📥 Generate Wage Slips")

    if st.button("📄 Generate Wage Slip Document", type="primary", use_container_width=True):
        with st.spinner("Generating DOCX wage slips..."):
            doc_buffer = create_wage_slip_doc(weeks_data)
            first_thu = weeks_data[0][0]
            last_wed = weeks_data[-1][1]
            filename = f"Wage_Slips_{first_thu.strftime('%d%b%Y')}_to_{last_wed.strftime('%d%b%Y')}.docx"
            st.session_state["doc_buffer"] = doc_buffer
            st.session_state["doc_filename"] = filename
            st.session_state["doc_slip_count"] = total_slips

    if "doc_buffer" in st.session_state and st.session_state["doc_buffer"]:
        st.download_button(
            label="⬇️ Download Wage Slips (.docx)",
            data=st.session_state["doc_buffer"],
            file_name=st.session_state["doc_filename"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.success(
            f"✅ Document ready: **{st.session_state['doc_filename']}** — "
            f"{st.session_state.get('doc_slip_count', '?')} wage slip(s)"
        )

    # ── Footer ──
    st.markdown("""
    <div class="footer-text">
        The Nandanvan Estate — Weekly Wage Slip Generator v1.0
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()

